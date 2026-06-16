import importlib.metadata
import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import pandas as pd


@dataclass
class ParsedTextBlock:
    text: str
    page_number: int | None = None
    section_title: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedTableBlock:
    rows: list[dict[str, Any]]
    markdown: str
    name: str | None = None
    page_number: int | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedAssetBlock:
    data: bytes
    extension: str
    asset_type: str = "image"
    page_number: int | None = None
    bbox: dict[str, Any] | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    text_blocks: list[ParsedTextBlock] = field(default_factory=list)
    tables: list[ParsedTableBlock] = field(default_factory=list)
    assets: list[ParsedAssetBlock] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)
    parser_name: str = "fallback"
    parser_version: str = "1"


class DocumentParser:
    DOCLING_EXTENSIONS = {".pdf", ".docx", ".pptx", ".html", ".md", ".png", ".jpg", ".jpeg"}

    def parse(self, path: Path) -> ParsedDocument:
        suffix = path.suffix.lower()
        if suffix in self.DOCLING_EXTENSIONS:
            try:
                return self._parse_with_docling(path)
            except ImportError:
                pass
            except Exception:
                if suffix in {".pdf", ".pptx"}:
                    raise
        if suffix == ".txt" or suffix == ".md":
            return ParsedDocument(
                text_blocks=[ParsedTextBlock(path.read_text(encoding="utf-8"))],
                metadata={"source": str(path)},
                parser_name="plain_text",
            )
        if suffix == ".csv":
            return self._parse_dataframe(pd.read_csv(path), path, "pandas_csv")
        if suffix in {".xlsx", ".xls"}:
            sheets = pd.read_excel(path, sheet_name=None)
            parsed = ParsedDocument(metadata={"source": str(path)}, parser_name="pandas_excel")
            for sheet_name, frame in sheets.items():
                table = self._dataframe_table(frame, name=str(sheet_name))
                parsed.tables.append(table)
            return parsed
        if suffix == ".docx":
            return self._parse_docx_fallback(path)
        if suffix in {".png", ".jpg", ".jpeg", ".webp"}:
            return ParsedDocument(
                text_blocks=[ParsedTextBlock("")],
                assets=[
                    ParsedAssetBlock(
                        data=path.read_bytes(),
                        extension=suffix,
                        asset_type="source_image",
                    )
                ],
                metadata={"source": str(path)},
                parser_name="image_asset",
            )
        raise ValueError(
            f"Unsupported file type {suffix!r}. Install the ingestion extra for Docling support."
        )

    def _parse_with_docling(self, path: Path) -> ParsedDocument:
        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.document_converter import DocumentConverter, PdfFormatOption
        from docling_core.types.doc import DocItemLabel

        pdf_options = PdfPipelineOptions(
            generate_picture_images=True,
            generate_page_images=False,
            do_picture_description=False,
            do_picture_classification=False,
        )
        converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pdf_options),
            }
        )
        result = converter.convert(path)
        document = result.document
        parsed = ParsedDocument(
            metadata={
                "source": str(path),
                "asset_extraction_errors": [],
                "tables_excluded_from_text": True,
            },
            parser_name="docling",
            parser_version=importlib.metadata.version("docling"),
        )
        text_labels = set(DocItemLabel) - {
            DocItemLabel.TABLE,
            DocItemLabel.DOCUMENT_INDEX,
        }
        page_numbers = sorted(getattr(document, "pages", {}))
        if page_numbers:
            for page_number in page_numbers:
                markdown = document.export_to_markdown(
                    page_no=page_number,
                    labels=text_labels,
                )
                if markdown.strip():
                    parsed.text_blocks.append(
                        ParsedTextBlock(
                            markdown,
                            page_number=page_number,
                            metadata={"content_source": "docling_page_without_tables"},
                        )
                    )
        else:
            markdown = document.export_to_markdown(labels=text_labels)
            if markdown.strip():
                parsed.text_blocks.append(
                    ParsedTextBlock(
                        markdown,
                        metadata={"content_source": "docling_document_without_tables"},
                    )
                )

        for index, table in enumerate(getattr(document, "tables", []), start=1):
            frame = table.export_to_dataframe(doc=document)
            page_number, bbox = self._provenance(table)
            parsed_table = self._dataframe_table(frame, name=f"table_{index}")
            parsed_table.page_number = page_number
            parsed_table.metadata = {
                "bbox": bbox,
                "source_ref": getattr(table, "self_ref", None),
            }
            parsed.tables.append(parsed_table)

        for picture in getattr(document, "pictures", []):
            try:
                image = picture.get_image(document)
                if image is None:
                    raise ValueError("Docling picture did not contain a generated image")
                buffer = io.BytesIO()
                image.save(buffer, format="PNG")
                page_number, bbox = self._provenance(picture)
                parsed.assets.append(
                    ParsedAssetBlock(
                        data=buffer.getvalue(),
                        extension=".png",
                        asset_type="document_image",
                        page_number=page_number,
                        bbox=bbox,
                        metadata={"source_ref": getattr(picture, "self_ref", None)},
                    )
                )
            except Exception as exc:
                parsed.metadata["asset_extraction_errors"].append(
                    {
                        "source_ref": getattr(picture, "self_ref", None),
                        "error": str(exc),
                    }
                )
        if path.suffix.lower() in {".png", ".jpg", ".jpeg"} and not parsed.assets:
            parsed.assets.append(
                ParsedAssetBlock(
                    data=path.read_bytes(),
                    extension=path.suffix,
                    asset_type="source_image",
                )
            )
        return parsed

    @staticmethod
    def _provenance(item: Any) -> tuple[int | None, dict[str, Any] | None]:
        provenance = getattr(item, "prov", None) or []
        if not provenance:
            return None, None
        first = provenance[0]
        bbox = getattr(first, "bbox", None)
        bbox_json = bbox.model_dump(mode="json") if bbox is not None else None
        return getattr(first, "page_no", None), bbox_json

    def _parse_dataframe(self, frame: pd.DataFrame, path: Path, parser_name: str) -> ParsedDocument:
        return ParsedDocument(
            tables=[self._dataframe_table(frame, name=path.stem)],
            metadata={"source": str(path)},
            parser_name=parser_name,
        )

    @staticmethod
    def _dataframe_table(frame: pd.DataFrame, name: str) -> ParsedTableBlock:
        clean = frame.fillna("")
        rows = clean.to_dict(orient="records")
        try:
            markdown = clean.to_markdown(index=False)
        except ImportError:
            markdown = clean.to_csv(index=False)
        return ParsedTableBlock(rows=rows, markdown=markdown, name=name)

    @staticmethod
    def _parse_docx_fallback(path: Path) -> ParsedDocument:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise ImportError("Install `uv sync --extra ingestion` to parse DOCX") from exc
        document = DocxDocument(path)
        parsed = ParsedDocument(metadata={"source": str(path)}, parser_name="python-docx")
        text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)
        if text:
            parsed.text_blocks.append(ParsedTextBlock(text))
        for index, table in enumerate(document.tables, start=1):
            if not table.rows:
                continue
            headers = [
                cell.text.strip() or f"column_{i}" for i, cell in enumerate(table.rows[0].cells)
            ]
            rows = [
                {headers[i]: cell.text.strip() for i, cell in enumerate(row.cells)}
                for row in table.rows[1:]
            ]
            frame = pd.DataFrame(rows)
            parsed.tables.append(DocumentParser._dataframe_table(frame, f"table_{index}"))
        return parsed
